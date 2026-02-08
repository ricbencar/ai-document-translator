"""
====================================================================================================
   AI DOCUMENT TRANSLATOR (ENGINEERING EDITION) - GUI
   Version: 1.0 (Final Production / Batch Support / Anti-Leak / Deep Context)
====================================================================================================

1. CORE METHODOLOGY: THE "DIGITAL TWIN"
   -------------------------------------------------------------------------------------------------
   This application operates as a Document Reconstruction Engine, not just a text translator.
   It treats .docx files as hierarchical XML structures to preserve strict visual fidelity.

   - ATOMIC DISSECTION: Uses `python-docx` to surgically parse the file into atomic XML units 
     (Paragraphs > Runs > Tables > Cells) rather than treating the document as a flat string.
   
   - DNA EXTRACTION: Before translation, the engine extracts the "Visual DNA" of every text segment.
     This includes Font Family, Size, RGB Color, Highlight, Bold/Italic flags, and critical 
     XML Numbering properties (`numPr`) to keep lists functional.
   
   - COMPILER INJECTION: Text is encapsulated with semantic tags (e.g., "{b}Text{/b}") and injected 
     into the local LLM. The model is forced into a deterministic "Compiler Mode" (Temp=0.0) 
     to process syntax without adding conversational filler.
   
   - HALLUCINATION SANITIZATION: The output passes through a multi-layer "Regex Firewall" to strip 
     AI meta-commentary (e.g., "Here is the translation"), Markdown artifacts, and broken tags.
   
   - SURGICAL RECONSTRUCTION: The document is rebuilt from zero. The sanitized translated text is 
     fused with the original "Visual DNA" and injected back into the XML skeleton, preserving 
     complex table layouts, headers, and footers exactly as they appeared in the source.

2. PRIMARY TRANSLATION ENGINES (Select One)
   -------------------------------------------------------------------------------------------------
   - SPECIALIST: TranslateGemma 3 (27B) - Optimized for high-fidelity technical specs.
     > Command: ollama pull translategemma:27b
   - POWERHOUSE: GPT-OSS 20B - Superior semantic understanding for legal/contractual specs.
     > Command: ollama pull gpt-oss:20b
   - POLYGLOT: Qwen-MT (Qwen 3) - Gold standard for CJK (Chinese/Japanese/Korean) manuals.
     > Command: ollama pull qwen-mt
   - SPEEDSTER: Qwen 3 (4B) - Ultra-fast; designed for legacy hardware and 8GB RAM systems.
     > Command: ollama pull qwen3:4b
   - LOCALIZER: T&L Specialist (50B) - Specialized for cultural adaptation and EIA reports.
     > Command: ollama pull ALIENTELLIGENCE/translationandlocalizationspecialist

3. SECONDARY CONSULTANT MODELS (Terminology Audit)
   -------------------------------------------------------------------------------------------------
   Used in "Consultant Mode" to double-check technical nomenclature without translating.
   - Engineering Manuals: `ollama pull ALIENTELLIGENCE/engineeringtechnicalmanuals`
   - Structural Engineer: `ollama pull joreilly86/structural_llama_3.0`
   - Civil Structure V2:  `ollama pull ALIENTELLIGENCE/civilstructureengineerv2`
   - Geotechnical:        `ollama pull ALIENTELLIGENCE/geotechnicalengineer`
   - Marine/Naval:        `ollama pull ALIENTELLIGENCE/marineengineernavalarchitect`
   - Mechanical V2:       `ollama pull ALIENTELLIGENCE/mechanicalengineer`
   - Electrical V2:       `ollama pull ALIENTELLIGENCE/electricalengineerv2`
   - Chemical Engineer:   `ollama pull ALIENTELLIGENCE/chemicalengineer`
   - Industrial Engineer: `ollama pull ALIENTELLIGENCE/industrialengineer`
   - Environmental Eng:   `ollama pull ALIENTELLIGENCE/environmentalengineer`
   - Env. Consulting:     `ollama pull ALIENTELLIGENCE/environmentalconsulting`

4. OPERATIONAL WORKFLOW
   -------------------------------------------------------------------------------------------------
   1. Initialize: Run `python translator_gui.py` and ensure the Ollama service is active.
   2. Ingest: Click "Browse" to select single or multiple .docx files for batching.
   3. Configure: Choose Source/Target languages and select your Primary Model.
   4. Execute: Click "Start Translation" and monitor progress via the real-time status log.
   5. Verify: Retrieve code-suffixed files (e.g., _DE.docx) from the source directory.

5. KEY BIBLIOGRAPHY
   -------------------------------------------------------------------------------------------------
   FOUNDATIONAL ARCHITECTURES
   - Jurafsky & Martin (2026): *Speech and Language Processing* (Transformer Foundations).
   - Tunstall et al. (2022): *NLP with Transformers* (Hugging Face / Practical Guide).
   - Raschka (2025): *Build a Large Language Model (From Scratch)* (Internal Mechanics).

   MACHINE TRANSLATION & LOCALIZATION
   - Koehn (2020): *Neural Machine Translation* (Core NMT Principles).
   - Kenny (2022): *Machine Translation for Everyone* (MT Literacy & Workflows).
   - Sun et al. (2025): *Translation Studies in the Age of AI* (Theory & Practice).

   ENGINEERING & PRODUCTION
   - Huyen (2025): *AI Engineering* (Production Systems & Scalability).
   - Iusztin & Labonne (2025): *The LLM Engineering Handbook* (RAG & Fine-tuning).
   - Bouchard & Peters (2025): *Building LLMs for Production* (Ops & Latency).

   ALIGNMENT & ETHICS
   - Christian (2020): *The Alignment Problem* (Safety & Value Alignment).
   - Narayanan & Kapoor (2024): *AI Snake Oil* (Critical capabilities analysis).
   
====================================================================================================
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import sys
import os
import time
import datetime
import re
import ollama
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy
import html
import json
from langdetect import detect, LangDetectException

# ====================================================================================================
#   CLASS: TRANSLATION ENGINE
#   The heart of the application. Runs on a background thread to keep the GUI responsive.
# ====================================================================================================

class TranslationEngine:
    """
    The engine encapsulates the specific 'Engineering Edition' logic for document reconstruction.
    It runs inside a separate thread to prevent the GUI from freezing.
    """
    def __init__(self, config, log_callback, progress_callback):
        self.config = config
        self.log = log_callback
        self.update_progress = progress_callback
        self.stop_flag = False
        self.translation_history = []
        
    def format_time(self, seconds):
        """Converts seconds into HH:MM:SS format."""
        return str(datetime.timedelta(seconds=int(seconds)))

    def update_history(self, new_text):
        """Updates the rolling window of context for the AI."""
        # Strip tags to keep context clean
        clean_text = re.sub(r'\{.*?\}', '', new_text)
        self.translation_history.append(clean_text)
        if len(self.translation_history) > self.config['context_window']:
            self.translation_history.pop(0)

    # --- 1. REFLEXION LOGIC (STABLE CORE) ---
    def _reflexion_check_is_chatty(self, text):
        """
        Asks the model itself if the output contains conversational filler.
        Returns True if the model detects 'chatter' (leaks).
        """
        # Optimization: Don't waste compute on very short segments
        if len(text.split()) < 5: return False
        
        try:
            response = ollama.chat(model=self.config['model_primary'], messages=[
                {'role': 'system', 'content': "CLASSIFIER TASK: Analyze the text. Return 'YES' if it contains conversational filler (e.g., 'Here is', 'Sure', 'I have translated', 'Note:', 'The table'). Return 'NO' if it contains ONLY the translated text."},
                {'role': 'user', 'content': f"Text: \"{text}\""}
            ], options={'num_predict': 5, 'temperature': 0.0}) 
            
            return "YES" in response['message']['content'].strip().upper()
        except:
            return False

    # --- 2. SANITIZATION FIREWALL (ADVANCED ANTI-LEAK) ---
    def _sanitize_output(self, text):
        """
        Aggressive post-processing to remove conversational filler, 
        fix encoding errors, strip HTML entities and repair malformed tags.
        """
        if not text: return ""

        # 1. Decode HTML Entities
        text = html.unescape(text)

        # 2. Fix Unicode Escapes & Broken Encoding (Fixes "\c3\ada" artifacts)
        try:
            if "\\u" in text:
                text = text.encode('utf-8').decode('unicode_escape')
            # Fix raw byte representations sometimes leaked by LLMs
            if "\\x" in text or "\\c3" in text:
                text = text.encode('latin1').decode('utf-8')
        except:
            pass

        # 3. REPAIR MALFORMED TAGS (Fixes {"b}, {\b}, {/b"}, etc.)
        text = re.sub(r'\{["\\]+([a-z]+)\}', r'{\1}', text)
        text = re.sub(r'\{/["\\]+([a-z]+)\}', r'{/\1}', text)
        text = re.sub(r'\{\{([a-z]+)\}\}', r'{\1}', text)

        # 4. STRIP LEAKED JSON ARTIFACTS (Fixes "1.1.1" prefixes)
        if '```' in text or 'json' in text or '"t":' in text:
            text = re.sub(r'```json', '', text)
            text = re.sub(r'```', '', text)
            # Remove isolated JSON keys if they leaked into text
            text = re.sub(r'\{"t":\s*', '', text)
            text = re.sub(r'^\s*"\s*', '', text)
            text = re.sub(r'\s*"\s*\}\s*$', '', text)

        # 5. Strip HTML-style tags (<br>, <b>, </div>)
        text = re.sub(r'<[^>]+>', '', text)

        # 6. Strip Markdown artifacts (**text**)
        text = text.replace('**', '').replace('__', '')

        # 7. Remove known placeholder leaks ("(tab)", "{tab}")
        leak_pattern = r'(?:\{|\()(/?(?:tab|br|nbsp|newline|b|i|u|s|ds|sup|sub|sc|caps|hl|user|id|var|empty))(?:\}|\))'
        text = re.sub(leak_pattern, ' ', text, flags=re.IGNORECASE)

        # 8. Strip "Introductory" Leaks (The "Firewall")
        start_leaks = [
            r'^The following table[:\.]?', 
            r'^Here is the translation[:\.]?',
            r'^Here are a few options[:\.]?',
            r'^Note[:\.]',
            r'^Translation[:\.]',
            r'^Translated text[:\.]',
            r'^Sure, here is[:\.]?',
            r'^This appears to be[:\.]?',
            r'^likely means[:\.]?'
        ]
        for pattern in start_leaks:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE).strip()

        # 9. Unicode Purge & Whitespace
        import unicodedata
        text = "".join(ch for ch in text if unicodedata.category(ch)[0] != "C" or ch in ['\n', '\t'])
        text = re.sub(r'\s+', ' ', text).strip()

        return text

    def _is_leak_detected(self, text, system_prompt_snippet=""):
        """
        Detects if the AI has outputted instructions instead of translation.
        """
        leak_triggers = [
            "You are a specialized", "TAG ARCHITECTURE", "technical constancy",
            "zero hallucination", "translate the input text", "standard & reference lock",
            "output only the translated", "context (previous text)", "strict rules",
            "maintain brevity", "translation:", "translated text:", "source:", "target:"
        ]
        if any(trigger in text.lower() for trigger in leak_triggers):
            return True
        if len(text) > 20 and system_prompt_snippet and system_prompt_snippet in text:
            return True
        return False

    # --- 3. FONT DNA & TAGGING ---

    def get_dominant_font_properties(self, paragraph):
        if not paragraph.runs: return None
        for run in paragraph.runs:
            if run.text.strip():
                return {
                    'name': run.font.name,
                    'size': run.font.size,
                    'color': run.font.color.rgb if run.font.color else None,
                    'highlight': run.font.highlight_color
                }
        return None

    def apply_font_properties(self, run, font_props):
        if not font_props: return
        if font_props['name']: run.font.name = font_props['name']
        if font_props['size']: run.font.size = font_props['size']
        if font_props['color']: run.font.color.rgb = font_props['color']
        if font_props['highlight']: run.font.highlight_color = font_props['highlight']

    def extract_tagged_text(self, paragraph):
        text_builder = ""
        for run in paragraph.runs:
            text = run.text
            if not text: continue
            
            text = text.replace('\x0b', '{br}').replace('\t', '{tab}')
            text = text.replace("{", "(").replace("}", ")")
            
            prefix = ""
            suffix = ""
            if run.bold: prefix += "{b}"; suffix = "{/b}" + suffix
            if run.italic: prefix += "{i}"; suffix = "{/i}" + suffix
            if run.underline: prefix += "{u}"; suffix = "{/u}" + suffix
            if run.font.strike: prefix += "{s}"; suffix = "{/s}" + suffix
            if run.font.double_strike: prefix += "{ds}"; suffix = "{/ds}" + suffix
            if run.font.superscript: prefix += "{sup}"; suffix = "{/sup}" + suffix
            if run.font.subscript: prefix += "{sub}"; suffix = "{/sub}" + suffix
            if run.font.small_caps: prefix += "{sc}"; suffix = "{/sc}" + suffix
            if run.font.all_caps: prefix += "{caps}"; suffix = "{/caps}" + suffix
            if run.font.highlight_color: prefix += "{hl}"; suffix = "{/hl}" + suffix
                
            text_builder += f"{prefix}{text}{suffix}"
        return text_builder

    def rebuild_paragraph(self, paragraph, translated_text, font_props, original_format):
        # 1. XML NUMBERING PROTECTION (Critical for Lists)
        p_element = paragraph._p
        original_numPr = None
        if p_element.pPr is not None and p_element.pPr.numPr is not None:
            original_numPr = deepcopy(p_element.pPr.numPr)

        # 2. CAPTURE LAYOUT
        p_style = paragraph.style
        p_alignment = paragraph.alignment
        p_left = paragraph.paragraph_format.left_indent
        p_right = paragraph.paragraph_format.right_indent
        p_first = paragraph.paragraph_format.first_line_indent
        p_before = paragraph.paragraph_format.space_before
        p_after = paragraph.paragraph_format.space_after
        p_tab_stops = []
        for ts in paragraph.paragraph_format.tab_stops:
            p_tab_stops.append((ts.position, ts.alignment, ts.leader))

        paragraph.clear()

        # 3. RESTORE ATTRIBUTES
        try:
            paragraph.style = p_style 
            paragraph.alignment = p_alignment
            paragraph.paragraph_format.left_indent = p_left
            paragraph.paragraph_format.right_indent = p_right
            paragraph.paragraph_format.first_line_indent = p_first
            paragraph.paragraph_format.space_before = p_before
            paragraph.paragraph_format.space_after = p_after
            for ts_data in p_tab_stops:
                paragraph.paragraph_format.tab_stops.add_tab_stop(*ts_data)
            
            # REINJECT NUMBERING
            if original_numPr is not None:
                if p_element.pPr is None: p_element.get_or_add_pPr()
                if p_element.pPr.numPr is not None: p_element.pPr.remove(p_element.pPr.numPr)
                p_element.pPr.insert(0, original_numPr)
        except Exception as e:
            self.log(f"[WARN] Format restoration partial fail: {e}")

        # 4. REBUILD CONTENT
        tag_pattern = r'(\{(?:b|i|u|s|ds|sup|sub|sc|caps|hl)\}.*?\{/(?:b|i|u|s|ds|sup|sub|sc|caps|hl)\}|\{br\}|\{tab\})'
        tokens = re.split(tag_pattern, translated_text)

        for token in tokens:
            if not token: continue
            
            if token == "{br}":
                paragraph.add_run().add_break()
                continue
            if token == "{tab}":
                paragraph.add_run().add_tab()
                continue

            is_bold = "{b}" in token
            is_italic = "{i}" in token
            is_underline = "{u}" in token
            is_strike = "{s}" in token
            is_dstrike = "{ds}" in token
            is_sup = "{sup}" in token
            is_sub = "{sub}" in token
            is_sc = "{sc}" in token
            is_caps = "{caps}" in token
            
            clean_text = re.sub(r'\{/?(?:b|i|u|s|ds|sup|sub|sc|caps|hl)\}', '', token)
            if not clean_text: continue

            run = paragraph.add_run()
            if font_props: self.apply_font_properties(run, font_props)

            if is_bold: run.bold = True
            if is_italic: run.italic = True
            if is_underline: run.underline = True
            if is_strike: run.font.strike = True
            if is_dstrike: run.font.double_strike = True
            if is_sup: run.font.superscript = True
            if is_sub: run.font.subscript = True
            if is_sc: run.font.small_caps = True
            if is_caps: run.font.all_caps = True
            
            if "{hl}" in token:
                if font_props and font_props.get('highlight'):
                    run.font.highlight_color = font_props['highlight']
                else:
                    run.font.highlight_color = 7
            run.text = clean_text

    # --- 4. CONSULTANT LOGIC ---

    def consult_expert_on_terminology(self, original_text, translated_text):
        if not self.config['model_secondary'] or self.config['model_secondary'] == "None":
            return translated_text
        if len(translated_text.split()) < 6: return translated_text

        system_prompt = f"""SYSTEM: TECHNICAL_AUDIT_ENGINE
TASK: Verify translation accuracy for Engineering Terms, Standards (ISO/DIN), and Units.
SOURCE_LANG: {self.config['lang_in']}
TARGET_LANG: {self.config['lang_out']}
OUTPUT PROTOCOL (JSON ONLY):
1. If the Draft is ACCURATE: return {{"status": "PASS"}}
2. If the Draft has ERRORS: return {{"status": "FIX", "corrected": "FIXED_STRING"}}
"""
        try:
            response = ollama.chat(model=self.config['model_secondary'], messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': f"SOURCE: \"{original_text}\"\nDRAFT: \"{translated_text}\""}
            ], options={'temperature': 0.0})
            
            raw_output = response['message']['content'].strip()
            
            # Robust Parsing for Consultant
            clean_json = raw_output
            if "```" in clean_json:
                clean_json = clean_json.split("```json")[-1].split("```")[0].strip()
            
            data = json.loads(clean_json)
            if data.get("status") == "FIX" and "corrected" in data:
                fixed_text = data["corrected"]
                if fixed_text and len(fixed_text) < (len(translated_text) * 2):
                    if not self._is_leak_detected(fixed_text):
                        return fixed_text
        except:
            pass
        return translated_text

    # --- 5. MAIN TRANSLATION BLOCK (ORCHESTRATOR) ---

    def translate_block_with_fidelity(self, paragraph, is_table_cell=False):
        if self.stop_flag: return False
        
        # Capture DNA
        t0 = time.time()
        font_props = self.get_dominant_font_properties(paragraph)
        original_format = paragraph.paragraph_format
        tagged_source = self.extract_tagged_text(paragraph)
        
        # Skip Logic
        clean_text = re.sub(r'\{.*?\}', '', tagged_source).strip()
        if not clean_text: return False
        if re.match(r'^[\d.]+$', clean_text): return False 
        
        # Logging (Aligned Format)
        context_str = "\n".join(self.translation_history)
        preview_in = (clean_text[:70] + '..') if len(clean_text) > 70 else clean_text
        self.log(f"┌── [SRC] \"{preview_in}\"")

        # Header/Prompt Setup
        word_count = len(clean_text.split())
        is_header = (clean_text.isupper() and word_count < 12) or (clean_text.istitle() and word_count < 4)
        header_instruction = "IMPORTANT: Input is a HEADLINE/TITLE. Translate it meaningfuly. It is NOT a brand name." if is_header else ""
        special_instruction = "CONTEXT NOTE: TABLE CELL." if is_table_cell else ""

        system_prompt = f"""SYSTEM MODE: DATA_PROCESSING_ENGINE 
OPCODE: TRANSLATE_ISO_STANDARD
SOURCE: {self.config['lang_in']}
TARGET: {self.config['lang_out']}

CRITICAL ENGINEERING CONSTRAINTS: 
1. OUTPUT FORMAT: STRICT JSON. Structure: {{"t": "translated_string"}}
2. TAG INVARIANCE: Input tags {{b}}, {{i}}, {{tab}} are MEMORY POINTERS. PRESERVE EXACTLY.
3. TERMINOLOGY LOCK: Do NOT translate Units or Standards (ISO, DIN, kg, mm).
4. HALLUCINATION FIREWALL: FORBIDDEN PHRASES: "Here is", "Note:", "Sure", "The table".

SPECIAL INSTRUCTIONS:
{header_instruction}
{special_instruction}

CONTEXT:
{context_str}

INPUT STREAM:"""

        try:
            # Primary Attempt
            response = ollama.chat(model=self.config['model_primary'], messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': f"\"{tagged_source}\""},
            ], options={'temperature': 0.0})
            
            raw_output = response['message']['content'].strip()
            result = raw_output
            parse_success = False
            
            # Robust Parsing Strategy (Regex First)
            json_match = re.search(r'\{.*"t"\s*:\s*"(.*?)".*\}', raw_output, re.DOTALL)
            if json_match:
                result = json_match.group(1).replace('\\"', '"')
                parse_success = True
            else:
                try:
                    clean_json = raw_output
                    if "```" in clean_json:
                        clean_json = clean_json.split("```json")[-1].split("```")[0].strip()
                    elif "'''" in clean_json:
                        clean_json = clean_json.split("'''json")[-1].split("'''")[0].strip()
                    
                    if clean_json.startswith("{"):
                        data = json.loads(clean_json)
                        if "t" in data:
                            result = data["t"]
                            parse_success = True
                except:
                    if result.startswith('"') and result.endswith('"'):
                        result = result[1:-1]

            # Sanitization (Firewall)
            result = self._sanitize_output(result)

            # Defensive Checks
            is_wrong_lang = False
            try:
                if len(result.split()) > 4:
                    detected = detect(result)
                    user_lang_prefix = self.config['lang_out'][:2].lower()
                    target_map = {'po': 'pt', 'en': 'en', 'sp': 'es', 'fr': 'fr', 'ge': 'de', 'ch': 'zh', 'it': 'it'}
                    target_iso = target_map.get(user_lang_prefix, 'xx') 
                    if target_iso != 'xx' and target_iso != detected and detected == 'en' and target_iso != 'en':
                        is_wrong_lang = True
            except LangDetectException: pass

            is_too_long = len(result) > (len(tagged_source) * 1.8) + 20
            
            # Reflexion Check
            is_chatty = False
            if not parse_success or is_too_long:
                 is_chatty = self._reflexion_check_is_chatty(result)
            
            is_explicit_leak = self._is_leak_detected(result, system_prompt)

            # Panic Mode
            if is_wrong_lang or is_too_long or is_chatty or is_explicit_leak:
                failure_reason = f"L:{is_wrong_lang} Len:{is_too_long} Ch:{is_chatty} Lk:{is_explicit_leak}"
                self.log(f"├── [WARN] Issue detected ({failure_reason}). Engaging Panic Mode...")
                
                retry_response = ollama.chat(model=self.config['model_primary'], messages=[
                    {'role': 'system', 'content': f"Translate from {self.config['lang_in']} to {self.config['lang_out']}."},
                    {'role': 'user', 'content': tagged_source}
                ], options={'temperature': 0.0})
                
                result = retry_response['message']['content'].strip()
                if result.startswith('"') and result.endswith('"'): result = result[1:-1]
                
                # Sanitize Panic Output
                result = self._sanitize_output(result)
                self.log(f"├── [RECOVERED] Fallback applied.")

            # Consultant Pass
            final_result = result
            if "{" not in result and len(clean_text.split()) > 5:
                consultant_result = self.consult_expert_on_terminology(clean_text, result)
                if consultant_result != result:
                    self.log(f"├── [CONSULTANT] Refined: \"{consultant_result[:40]}...\"")
                    final_result = consultant_result

            # Finalize
            self.rebuild_paragraph(paragraph, final_result, font_props, original_format)
            clean_translation = re.sub(r'\{.*?\}', '', final_result)
            self.update_history(clean_translation)
            
            elapsed = time.time() - t0
            preview_out = (clean_translation[:70] + '..') if len(clean_translation) > 70 else clean_translation
            
            # Aligned Log Output with Timer at End
            self.log(f"└── [TRG] \"{preview_out}\"   [{elapsed:.2f}s]")
            
            return True

        except Exception as e:
            self.log(f"└── [ERROR] Block Failed: {str(e)}")
            return False

    # --- 6. BATCH PROCESSING LOOP ---

    def run(self):
        input_files = self.config.get('file_paths', [])
        if not input_files and 'file_path' in self.config:
             input_files = [self.config['file_path']]

        if not input_files:
            self.log("[CRITICAL] No files provided.")
            return

        self.log(f"-> Batch Started: {len(input_files)} files queued.")
        if self.config['model_secondary'] and self.config['model_secondary'] != "None":
             self.log(f"-> Consultant Mode: ACTIVE ({self.config['model_secondary']})")

        # Global Pre-Scan
        self.log("-> Scanning all files to calculate total workload...")
        total_global_items = 0
        for idx, f_path in enumerate(input_files):
            if self.stop_flag: return
            try:
                temp_doc = Document(f_path)
                for section in temp_doc.sections:
                    for part in [section.header, section.footer]:
                        total_global_items += len(part.paragraphs)
                        for table in part.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    total_global_items += sum(1 for p in cell.paragraphs if p.text.strip())
                total_global_items += len(temp_doc.paragraphs)
                for table in temp_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            total_global_items += sum(1 for p in cell.paragraphs if p.text.strip())
            except Exception as e:
                self.log(f"[WARN] Skipping scan for {os.path.basename(f_path)}: {e}")

        self.log(f"-> Total Workload: {total_global_items} blocks across {len(input_files)} files.")

        # Processing Loop
        start_time_global = time.time()
        global_processed_count = 0

        for file_index, input_path in enumerate(input_files):
            if self.stop_flag: break
            
            filename_base = os.path.basename(input_path)
            path_parts = os.path.splitext(input_path)
            output_path = f"{path_parts[0]}_translated{path_parts[1]}"
            
            self.log(f"------------------------------------------------")
            self.log(f"-> PROCESSING FILE {file_index + 1}/{len(input_files)}: {filename_base}")
            
            try:
                doc = Document(input_path)
            except Exception as e:
                self.log(f"[ERROR] Failed to load {filename_base}: {e}")
                continue

            processed_parts = set()

            # Headers & Footers
            for section in doc.sections:
                for part in [section.header, section.footer]:
                    if part.part not in processed_parts:
                        processed_parts.add(part.part)
                        for para in part.paragraphs:
                            if self.stop_flag: break
                            global_processed_count += 1
                            self.update_progress(global_processed_count, total_global_items, f"[{file_index+1}/{len(input_files)}] Header/Footer...")
                            self.translate_block_with_fidelity(para)
                        for table in part.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        if not para.text.strip(): continue
                                        if self.stop_flag: break
                                        global_processed_count += 1
                                        self.update_progress(global_processed_count, total_global_items, f"[{file_index+1}/{len(input_files)}] H/F Table...")
                                        self.translate_block_with_fidelity(para, is_table_cell=True)

            # Body Paragraphs
            for para in doc.paragraphs:
                if self.stop_flag: break
                global_processed_count += 1
                self.update_progress(global_processed_count, total_global_items, f"[{file_index+1}/{len(input_files)}] Body...")
                self.translate_block_with_fidelity(para)

            # Body Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if not para.text.strip(): continue
                            if self.stop_flag: break
                            global_processed_count += 1
                            self.update_progress(global_processed_count, total_global_items, f"[{file_index+1}/{len(input_files)}] Tables...")
                            self.translate_block_with_fidelity(para, is_table_cell=True)

            if not self.stop_flag:
                try:
                    doc.save(output_path)
                    self.log(f"-> SAVED: {os.path.basename(output_path)}")
                except Exception as e:
                    self.log(f"[ERROR] Could not save {filename_base}: {e}")

        if not self.stop_flag:
            total_time = time.time() - start_time_global
            self.log(f"================================================")
            self.log(f"-> BATCH COMPLETE. Time: {self.format_time(total_time)}")
            self.update_progress(total_global_items, total_global_items, "Batch Done.")
        else:
            self.log("-> ABORTED BY USER.")


# ====================================================================================================
#   CLASS: TRANSLATOR APP (GUI)
# ====================================================================================================

class TranslatorApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Offline AI Translator - Engineering Edition")
        self.geometry("900x750")
        self.configure(bg="#2b2b2b")
        style = ttk.Style()
        style.theme_use('clam')
        
        # Colors
        self.bg_color = "#2b2b2b"
        self.fg_color = "#ffffff"
        self.accent_color = "#4a90e2"
        
        style.configure("TFrame", background=self.bg_color)
        style.configure("TLabel", background=self.bg_color, foreground=self.fg_color, font=("Segoe UI", 10))
        style.configure("Header.TLabel", font=("Segoe UI", 16, "bold"), foreground=self.accent_color)
        style.configure("TButton", background="#3c3c3c", foreground="white", borderwidth=0, focuscolor="none")
        style.map("TButton", background=[("active", self.accent_color)])
        style.configure("TProgressbar", troughcolor="#3c3c3c", background=self.accent_color, thickness=20)
        style.configure("TEntry", fieldbackground="#3c3c3c", foreground="white")
        style.configure("TLabelframe", background=self.bg_color, foreground=self.fg_color)
        style.configure("TLabelframe.Label", background=self.bg_color, foreground=self.accent_color, font=("Segoe UI", 10, "bold"))

        # State Variables
        self.file_path = tk.StringVar()
        self.lang_in = tk.StringVar(value="Portuguese (Portugal)")  # Default Input
        self.lang_out = tk.StringVar(value="English (UK)")          # Default Output
        self.model_primary = tk.StringVar(value="") 
        self.model_secondary = tk.StringVar(value="ALIENTELLIGENCE/engineeringtechnicalmanuals")
        self.temperature = tk.DoubleVar(value=0.0)
        self.progress_var = tk.DoubleVar(value=0)
        self.status_var = tk.StringVar(value="Idle - Ready to load document")
        self.engine = None
        self.worker_thread = None

        self.create_widgets()

    def get_installed_models(self):
        """
        Robustly fetches the list of models using a Hybrid Approach (API + CLI).
        Returns a sorted list of model names or an empty list if none found.
        """
        models = set()

        # --- Method 1: Python Library (Fastest) ---
        try:
            response = ollama.list()
            
            # CASE A: New Library Version (Returns Object with .models attribute)
            if hasattr(response, 'models'):
                for m in response.models:
                    # Some versions use .model, others .name
                    if hasattr(m, 'model'):
                        models.add(m.model)
                    elif hasattr(m, 'name'):
                        models.add(m.name)
                        
            # CASE B: Older Library Version (Returns Dictionary)
            elif isinstance(response, dict) and 'models' in response:
                for m in response['models']:
                    # Dict keys can be 'model' or 'name'
                    models.add(m.get('model') or m.get('name'))
                    
        except Exception as e:
            print(f"Ollama Library fetch failed: {e}")

        # --- Method 2: CLI Backup (If Library fails) ---
        # Only run this if Method 1 returned nothing or crashed
        if not models:
            try:
                # Windows specific flag to hide the command window
                startup_info = None
                if sys.platform == "win32":
                    startup_info = subprocess.STARTUPINFO()
                    startup_info.dwFlags |= subprocess.STARTF_USESHOWWINDOW

                result = subprocess.run(
                    ['ollama', 'list'], 
                    capture_output=True, text=True, encoding='utf-8', errors='ignore',
                    startupinfo=startup_info
                )
                
                if result.returncode == 0:
                    lines = result.stdout.strip().splitlines()
                    # Skip the header row ("NAME ID SIZE...")
                    for line in lines[1:]:
                        if line.strip():
                            parts = line.split()
                            if parts:
                                models.add(parts[0]) # First column is always the name
            except Exception as e:
                print(f"CLI fetch failed: {e}")

        # Return sorted list or empty list (NO fake fallbacks)
        return sorted(list(models))

    def create_widgets(self):
        main_pad = 20
        # Header
        header = ttk.Frame(self)
        header.pack(fill="x", padx=main_pad, pady=(20, 10))
        ttk.Label(header, text="AI DOCUMENT TRANSLATOR", style="Header.TLabel").pack(side="left")
        ttk.Label(header, text="  |  Engineering Documents", foreground="#888").pack(side="left", pady=(8,0))

        # File
        f_frame = ttk.LabelFrame(self, text="INPUT DOCUMENT(S) (.docx)", padding=15)
        f_frame.pack(fill="x", padx=main_pad, pady=5)
        ttk.Entry(f_frame, textvariable=self.file_path, font=("Consolas", 10)).pack(side="left", fill="x", expand=True, padx=(0,10))
        ttk.Button(f_frame, text="BROWSE...", command=self.browse_file, width=15).pack(side="right")

        # Config
        c_frame = ttk.LabelFrame(self, text="CONFIGURATION", padding=15)
        c_frame.pack(fill="x", padx=main_pad, pady=5)
        
        # --- LANGUAGE SELECTION ---
        language_options = [
            "Portuguese (Portugal)", "Portuguese (Brazil)", "English (UK)",       "English (US)",
            "Spanish",               "French",              "German",             "Italian",
            "Chinese (Mandarin)",    "Arabic",              "Japanese",           "Korean",
            "Czech",                 "Danish",              "Dutch",              "Finnish",
            "Greek",                 "Hindi",               "Hungarian",          "Indonesian",
            "Norwegian",             "Polish",              "Romanian",           "Russian",
            "Swedish",               "Thai",                "Turkish",            "Vietnamese"
        ]

        ttk.Label(c_frame, text="Source:").grid(row=0, column=0, sticky="w", padx=5)
        self.cbox_lang_in = ttk.Combobox(c_frame, textvariable=self.lang_in, values=language_options, width=23, state="readonly")
        self.cbox_lang_in.grid(row=0, column=1, padx=5)
        
        ttk.Label(c_frame, text="Target:").grid(row=0, column=2, sticky="w", padx=(20,5))
        self.cbox_lang_out = ttk.Combobox(c_frame, textvariable=self.lang_out, values=language_options, width=23, state="readonly")
        self.cbox_lang_out.grid(row=0, column=3, padx=5)
        
        # --- DYNAMIC MODEL LOADING ---
        available_models = self.get_installed_models()
        
        # 1. Validate Primary Model (Smart Priority)
        # If the currently set primary model isn't installed, try to find a smart default
        if self.model_primary.get() not in available_models:
            if any("translategemma" in m and "27b" in m for m in available_models):
                match = next((m for m in available_models if "translategemma" in m and "27b" in m), "")
                self.model_primary.set(match)
            elif any("translategemma" in m and "12b" in m for m in available_models):
                match = next((m for m in available_models if "translategemma" in m and "12b" in m), "")
                self.model_primary.set(match)
            elif any("gpt-oss:20b" in m for m in available_models):
                match = next((m for m in available_models if "gpt-oss:20b" in m), "")
                self.model_primary.set(match)
            elif available_models:
                self.model_primary.set(available_models[0])
            else:
                self.model_primary.set("")

        ttk.Label(c_frame, text="Primary Model:").grid(row=1, column=0, sticky="w", padx=5, pady=10)
        cbox_m = ttk.Combobox(c_frame, textvariable=self.model_primary, width=23, state="readonly")
        cbox_m['values'] = available_models
        cbox_m.grid(row=1, column=1, padx=5)
        
        # 2. Validate Consultant Model (Strict Check)
        # If the default hardcoded consultant isn't installed, force it to "None"
        current_secondary = self.model_secondary.get()
        
        # Check for exact match OR match with ':latest' tag
        is_secondary_installed = (
            current_secondary in available_models or 
            f"{current_secondary}:latest" in available_models
        )
        
        if not is_secondary_installed:
            self.model_secondary.set("None")

        ttk.Label(c_frame, text="Consultant:").grid(row=1, column=2, sticky="w", padx=(20,5))
        cbox_c = ttk.Combobox(c_frame, textvariable=self.model_secondary, width=23, state="readonly")
        cbox_c['values'] = ["None"] + available_models
        cbox_c.grid(row=1, column=3, padx=5)

        # --- TEMPERATURE SECTION ---
        ttk.Label(c_frame, text="Temp:").grid(row=2, column=0, sticky="w", padx=5)
        
        # 1. Create a specific variable for the formatted text
        self.temp_display = tk.StringVar(value=f"{self.temperature.get():.2f}")

        # 2. Define a callback to formatting the number as the slider moves
        def update_temp_label(val):
            self.temp_display.set(f"{float(val):.2f}")

        # 3. Add the 'command' argument to the Scale and bind the Label to the formatted string
        ttk.Scale(c_frame, variable=self.temperature, from_=0.0, to=1.0, command=update_temp_label).grid(row=2, column=1, sticky="ew", padx=5)
        ttk.Label(c_frame, textvariable=self.temp_display, width=4).grid(row=2, column=2, sticky="w")

        # Actions
        a_frame = ttk.Frame(self)
        a_frame.pack(fill="x", padx=main_pad, pady=15)
        self.btn_start = ttk.Button(a_frame, text="START TRANSLATION", command=self.start_translation)
        self.btn_start.pack(side="left", fill="x", expand=True, ipady=10, padx=(0,10))
        self.btn_stop = ttk.Button(a_frame, text="STOP", command=self.stop_translation, state="disabled", width=15)
        self.btn_stop.pack(side="right", fill="x", ipady=10)

        # Progress & Log
        p_frame = ttk.Frame(self)
        p_frame.pack(fill="x", padx=main_pad, pady=5)
        ttk.Label(p_frame, textvariable=self.status_var, foreground=self.accent_color).pack(anchor="w")
        ttk.Progressbar(p_frame, variable=self.progress_var).pack(fill="x", ipady=2)
        
        l_frame = ttk.LabelFrame(self, text="LOG", padding=10)
        l_frame.pack(fill="both", expand=True, padx=main_pad, pady=(5,20))
        
        self.txt_console = scrolledtext.ScrolledText(l_frame, height=8, bg="#1e1e1e", fg="#00ff00", font=("Consolas", 9), borderwidth=0)
        self.txt_console.pack(fill="both", expand=True)

    def browse_file(self):
        files = filedialog.askopenfilenames(filetypes=[("Word", "*.docx")])
        if files:
            self.selected_files = files
            if len(files) == 1:
                self.file_path.set(files[0])
            else:
                self.file_path.set(f"({len(files)} files selected) - {files[0]}...")
            
            # Reset UI
            self.txt_console.delete(1.0, tk.END)
            self.progress_var.set(0)
            self.status_var.set("Ready to start.")
            self.log_message(f"-> Loaded {len(files)} file(s). Configured and ready.")

    def start_translation(self):
        target_files = []
        if hasattr(self, 'selected_files') and self.selected_files:
            target_files = self.selected_files
        elif self.file_path.get():
            target_files = [self.file_path.get().replace('"', '')]

        if not target_files:
            messagebox.showerror("Error", "Please select valid DOCX files.")
            return

        if not self.model_primary.get():
            messagebox.showerror("Error", "Please select a Primary Translation Model.")
            return

        sec_model = self.model_secondary.get()
        if sec_model and sec_model != "None":
            installed = self.get_installed_models()
            is_installed = (
                sec_model in installed or 
                f"{sec_model}:latest" in installed or 
                sec_model.replace(":latest", "") in installed
            )
            if not is_installed:
                if not messagebox.askyesno("Model Warning", f"Consultant '{sec_model}' not installed. Download now?"):
                    return

        config = {
            'file_paths': target_files,
            'lang_in': self.lang_in.get(),
            'lang_out': self.lang_out.get(),
            'model_primary': self.model_primary.get(),
            'model_secondary': self.model_secondary.get(),
            'temperature': self.temperature.get(),
            'context_window': 10
        }

        self.txt_console.delete(1.0, tk.END)
        self.toggle_ui_state(True)
        self.engine = TranslationEngine(config, self.log_message, self.update_progress_bar)
        self.worker_thread = threading.Thread(target=self.run_worker)
        self.worker_thread.start()

    def log_message(self, msg):
        self.after(0, lambda: self.txt_console.insert(tk.END, f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {msg}\n"))
        self.after(0, lambda: self.txt_console.see(tk.END))

    def update_progress_bar(self, current, total, msg):
        percent = min((current / total) * 100, 100) if total > 0 else 0
        self.after(0, lambda: self.progress_var.set(percent))
        self.after(0, lambda: self.status_var.set(f"{msg} ({int(percent)}%)"))

    def toggle_ui_state(self, is_running):
        state = "disabled" if is_running else "normal"
        self.btn_start.config(state=state)
        self.btn_stop.config(state="normal" if is_running else "disabled")

    def run_worker(self):
        try:
            self.engine.run()
        except Exception as e:
            self.log_message(f"FATAL ERROR: {e}")
        finally:
            self.after(0, lambda: self.toggle_ui_state(False))

    def stop_translation(self):
        if self.engine:
            self.log_message("Stopping translation... please wait.")
            self.engine.stop_flag = True

if __name__ == "__main__":
    app = TranslatorApp()
    app.mainloop()